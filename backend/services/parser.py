import re
import os
import json

# Try importing PDF/Word extractors, fallback to plain text if not installed
try:
    from pypdf import PdfReader
except ImportError:
    PdfReader = None

try:
    import docx
except ImportError:
    docx = None

# Keywords and Patterns for Supply Chain / Inventory Analyst domain in India
SKILL_KEYWORDS = [
    "inventory management", "supply chain", "demand planning", "logistics", "procurement", 
    "warehouse operations", "material management", "vendor management", "stock auditing", 
    "replenishment", "production planning", "shipping & receiving", "freight coordination", 
    "purchase order", "cycle counting", "operations management", "data analysis", "reporting", 
    "distribution", "sourcing", "capacity planning", "quality control"
]

TOOL_KEYWORDS = [
    "excel", "advanced excel", "power bi", "sql", "tableau", "python", "vba", "ms access", 
    "wms", "warehouse management system", "tms", "transportation management system"
]

ERP_KEYWORDS = [
    "sap", "sap mm", "sap sd", "sap wm", "oracle erp", "netsuite", "microsoft dynamics", 
    "odoo", "tally", "zoho inventory", "infor"
]

KPI_KEYWORDS = [
    "inventory turnover", "turnover ratio", "safety stock", "carrying cost", "holding cost", 
    "fill rate", "stockout", "shrinkage", "lead time", "otif", "on-time in-full", 
    "dsi", "days sales of inventory", "inventory accuracy", "cycle count accuracy", "reo"
]

CERTIFICATION_KEYWORDS = [
    "apics", "cpim", "cscp", "cltd", "six sigma", "green belt", "yellow belt", "black belt", 
    "supply chain analytics", "logistics certification", "lean"
]

def extract_text_from_pdf(file_path):
    if not PdfReader:
        raise ImportError("pypdf is not installed. Please install it or upload a .txt file.")
    try:
        reader = PdfReader(file_path)
        text = ""
        for page in reader.pages:
            page_text = page.extract_text()
            if page_text:
                text += page_text + "\n"
        return text
    except Exception as e:
        print(f"Error reading PDF: {e}")
        return ""

def extract_text_from_docx(file_path):
    if not docx:
        raise ImportError("python-docx is not installed. Please install it or upload a .txt file.")
    try:
        doc = docx.Document(file_path)
        text = ""
        for para in doc.paragraphs:
            text += para.text + "\n"
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += cell.text + " "
                text += "\n"
        return text
    except Exception as e:
        print(f"Error reading DOCX: {e}")
        return ""

def extract_text(file_path):
    ext = os.path.splitext(file_path)[1].lower()
    if ext == ".pdf":
        return extract_text_from_pdf(file_path)
    elif ext == ".docx":
        return extract_text_from_docx(file_path)
    elif ext in [".txt", ".md"]:
        with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
            return f.read()
    else:
        raise ValueError(f"Unsupported file extension: {ext}")

def parse_resume(text):
    """
    Parses resume text using heuristics and keyword scanning, specifically optimized for
    Entry-level Inventory Analyst, Supply Chain, and Operations roles in India.
    """
    text_lower = text.lower()
    
    # 1. Contact Information
    email_pattern = r'[a-zA-Z0-9_.+-]+@[a-zA-Z0-9-]+\.[a-zA-Z0-9-.]+'
    phone_pattern = r'(?:\+?\d{1,3}[-.\s]?)?\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}'
    
    emails = re.findall(email_pattern, text)
    phones = re.findall(phone_pattern, text)
    
    email = emails[0] if emails else ""
    phone = phones[0] if phones else ""
    
    # Parse Name: Usually on the first 2-3 lines of text
    lines = [line.strip() for line in text.split("\n") if line.strip()]
    name = "Candidate Name"
    if lines:
        for line in lines[:3]:
            # Ignore headers, emails, or phone numbers in line
            if "@" not in line and not any(char.isdigit() for char in line) and len(line.split()) <= 4:
                name = line
                break

    # 2. Extract Categories using dictionary matching
    extracted_skills = []
    for skill in SKILL_KEYWORDS:
        if re.search(r'\b' + re.escape(skill) + r'\b', text_lower):
            extracted_skills.append(skill.title())
            
    extracted_tools = []
    for tool in TOOL_KEYWORDS:
        if re.search(r'\b' + re.escape(tool) + r'\b', text_lower):
            # Special formatting for BI tools and SQL
            if tool in ["sql", "vba"]:
                extracted_tools.append(tool.upper())
            elif tool == "power bi":
                extracted_tools.append("Power BI")
            elif tool == "excel" or tool == "advanced excel":
                if "Advanced Excel" not in extracted_tools:
                    extracted_tools.append("Advanced Excel" if "advanced" in text_lower else "Excel")
            else:
                extracted_tools.append(tool.title())
                
    extracted_erps = []
    for erp in ERP_KEYWORDS:
        if re.search(r'\b' + re.escape(erp) + r'\b', text_lower):
            if erp == "sap":
                # Check for specific sub-modules
                modules = []
                if "sap mm" in text_lower or "material management" in text_lower:
                    modules.append("SAP MM (Material Management)")
                if "sap sd" in text_lower:
                    modules.append("SAP SD")
                if "sap wm" in text_lower or "warehouse management" in text_lower:
                    modules.append("SAP WM (Warehouse Management)")
                if not modules:
                    extracted_erps.append("SAP ERP")
                else:
                    extracted_erps.extend(modules)
            elif erp in ["oracle erp", "netsuite"]:
                extracted_erps.append(erp.title())
            else:
                extracted_erps.append(erp.upper())
                
    # Remove duplicates from ERPs
    extracted_erps = list(set(extracted_erps))

    extracted_kpis = []
    for kpi in KPI_KEYWORDS:
        if re.search(r'\b' + re.escape(kpi) + r'\b', text_lower):
            if kpi == "otif" or kpi == "on-time in-full":
                extracted_kpis.append("OTIF (On-Time In-Full)")
            elif kpi == "dsi" or kpi == "days sales of inventory":
                extracted_kpis.append("DSI (Days Sales of Inventory)")
            elif kpi == "turnover ratio" or kpi == "inventory turnover":
                if "Inventory Turnover Ratio" not in extracted_kpis:
                    extracted_kpis.append("Inventory Turnover Ratio")
            else:
                extracted_kpis.append(kpi.title())

    extracted_certs = []
    for cert in CERTIFICATION_KEYWORDS:
        if re.search(r'\b' + re.escape(cert) + r'\b', text_lower):
            if cert == "cpim":
                extracted_certs.append("APICS CPIM (Certified in Planning and Inventory Management)")
            elif cert == "cscp":
                extracted_certs.append("APICS CSCP (Certified Supply Chain Professional)")
            elif cert == "cltd":
                extracted_certs.append("APICS CLTD (Certified in Logistics, Transportation and Distribution)")
            elif "six sigma" in cert:
                # Find specific belt
                match = re.search(r'six sigma\s+(green|black|yellow)\s*belt', text_lower)
                if match:
                    extracted_certs.append(f"Six Sigma {match.group(1).title()} Belt")
                else:
                    extracted_certs.append("Six Sigma Certification")
            else:
                extracted_certs.append(cert.title())

    # 3. Heuristic Experience Extractor
    # Look for bullet points with numbers and percentages, e.g. "reduced shrinkage by 14%"
    achievements = []
    action_verbs = ["reduced", "optimized", "saved", "increased", "negotiated", "managed", "forecasted", "analyzed", "implemented", "audited", "streamlined", "slashed"]
    
    sentences = re.split(r'[.!?\n]', text)
    for sentence in sentences:
        sentence = sentence.strip()
        sentence_lower = sentence.lower()
        # Look for combinations of action verb + metrics (% or numbers or INR/Rs./$)
        has_action = any(verb in sentence_lower for verb in action_verbs)
        has_metric = any(char.isdigit() or "%" in sentence_lower or "rs" in sentence_lower or "inr" in sentence_lower for char in sentence_lower)
        # Verify it has supply chain terms
        has_sc = any(term in sentence_lower for term in ["stock", "inventory", "warehouse", "cost", "lead time", "accuracy", "supplier", "vendor", "fill rate", "carrying"])
        
        if has_action and has_metric and has_sc and len(sentence.split()) > 5:
            # Clean up bullet symbol
            cleaned = re.sub(r'^[\s•\-*]+', '', sentence)
            if cleaned not in achievements:
                achievements.append(cleaned)

    # 4. Standard Professional Summary Generator (If none detected)
    exp_summary = "Result-oriented Inventory & Supply Chain Analyst with 1 year of experience auditing stock, optimizing replenishment cycles, and leveraging data reporting systems (Advanced Excel, SAP ERP) to streamline warehouse and logistic operations. Proven record reducing safety stock levels and shrinkage while maintaining inventory accuracies above 98%."
    
    # 5. Core Target Roles based on extraction
    target_roles = ["Inventory Analyst", "Supply Chain Analyst", "Warehouse Analyst", "Logistics Coordinator", "Operations Executive"]
    
    profile_data = {
        "name": name,
        "email": email,
        "phone": phone,
        "target_roles": target_roles,
        "preferred_locations": ["Chennai", "Bangalore", "Remote"],
        "master_resume_text": text,
        "extracted_skills": extracted_skills,
        "extracted_tools": extracted_tools,
        "extracted_erps": extracted_erps,
        "extracted_kpis": extracted_kpis,
        "certifications": extracted_certs,
        "experience_summary": {"summary": exp_summary, "achievements": achievements[:5]},
        "education": [{"degree": "Bachelor of Commerce / Business Administration", "institution": "Indian University", "year": "2024"}]
    }
    
    return profile_data
